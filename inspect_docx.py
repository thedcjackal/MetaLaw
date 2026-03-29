import docx
import os

doc_path = r"c:\MetaLaw\translated_document.docx"
if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    print("\n--- BEGIN DOCUMENT TEXT ---\n")
    print("\n".join(full_text))
    print("\n--- END DOCUMENT TEXT ---\n")
else:
    print(f"Error: {doc_path} not found.")
